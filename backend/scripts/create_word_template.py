"""
Word 보고서 템플릿 생성 스크립트
python-docx를 사용하여 {{플레이스홀더}} 기반 Word 문서 템플릿을 생성합니다.
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


FONT_NAME = '맑은 고딕'


def set_korean_font(run, font_name=FONT_NAME):
    """run에 한글 폰트를 올바르게 설정"""
    run.font.name = font_name
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def set_style_korean_font(style, font_name=FONT_NAME):
    """스타일에 한글 폰트를 올바르게 설정"""
    style.font.name = font_name
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = style.element.makeelement(qn('w:rPr'), {})
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def create_project_report_template():
    """과제 보고서 Word 템플릿 생성"""
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.size = Pt(10)
    set_style_korean_font(style)

    # Heading 스타일에도 한글 폰트 설정
    for i in range(5):
        try:
            hs = doc.styles[f'Heading {i}']
            set_style_korean_font(hs)
        except KeyError:
            pass

    # ─── 제목 ───
    title = doc.add_heading(level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('{{과제명}}')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)
    set_korean_font(run)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('{{과제년도}}년 과제 보고서')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    set_korean_font(run)

    doc.add_paragraph()  # 여백

    # ─── 기본 정보 테이블 ───
    doc.add_heading('1. 과제 기본 정보', level=1)

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ('사업부', '{{사업부}}', '프로세스', '{{프로세스}}'),
        ('과제구분', '{{과제구분}}', '과제영역', '{{과제영역}}'),
        ('과제 PL', '{{과제PL}}', '작성자', '{{작성자}}'),
        ('시작월', '{{시작}}월', '종료월', '{{종료}}월'),
        ('진행상태', '{{진행상태}}', '진행률', '{{진행률}}%'),
        ('PoC 과제', '{{PoC과제여부}}', '중점 과제', '{{중점과제여부}}'),
    ]

    for row_idx, (l1, v1, l2, v2) in enumerate(info_data):
        row = table.rows[row_idx]
        for col_idx, text in enumerate([l1, v1, l2, v2]):
            cell = row.cells[col_idx]
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    set_korean_font(run)
                    if col_idx % 2 == 0:  # 라벨 열
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()

    # ─── 과제 상세 설명 ───
    doc.add_heading('2. 과제 상세 설명', level=1)
    doc.add_paragraph('{{과제상세설명}}')

    doc.add_paragraph()

    # ─── 성과 목록 ───
    doc.add_heading('3. 연결 성과 항목', level=1)

    perf_table = doc.add_table(rows=6, cols=6)
    perf_table.style = 'Table Grid'
    perf_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    headers = ['No', '성과항목', '대분류', '단위', '목표수준', '실적수준']
    for col_idx, h in enumerate(headers):
        cell = perf_table.rows[0].cells[col_idx]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(8)
                set_korean_font(run)

    # 성과 데이터 행 (5개)
    for i in range(1, 6):
        row = perf_table.rows[i]
        row.cells[0].text = str(i)
        row.cells[1].text = f'{{{{성과{i}_성과항목}}}}'
        row.cells[2].text = f'{{{{성과{i}_대분류}}}}'
        row.cells[3].text = f'{{{{성과{i}_단위}}}}'
        row.cells[4].text = f'{{{{성과{i}_목표수준}}}}'
        row.cells[5].text = f'{{{{성과{i}_실적수준}}}}'
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8)
                    set_korean_font(run)

    doc.add_paragraph()

    # ─── 비고 ───
    doc.add_heading('4. 비고', level=1)
    doc.add_paragraph('이 문서는 자동 생성된 과제 보고서 템플릿입니다.')

    # 저장
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'templates', 'document_type_word')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, 'project_report_template.docx')
    doc.save(output_path)
    print(f'Word 템플릿 생성 완료: {output_path}')


if __name__ == '__main__':
    create_project_report_template()
