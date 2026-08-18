"""
기획서를 Word 로 쓴다.

⚠️ **여기서 내용을 만들지 않는다.** document.assemble() 이 조립한 블록을 그대로
   옮길 뿐이다. 화면 미리보기와 이 파일이 같은 블록을 읽으므로, 화면에서 본 것과
   내보낸 문서가 어긋나지 않는다. 여기서 "문서니까 이건 빼자" 같은 판단을 하면
   그 순간 둘이 갈라진다.

⚠️ **빼지 않는 것들이 있다.** 비어 있는 구간, 합성 데이터 경고, 안 매긴 것 —
   문서로는 지저분해 보이지만 그게 이 모듈이 지켜 온 것이다. 깔끔한 문서가
   목적이면 애초에 조립할 이유가 없다.
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# 한글 글꼴. 지정하지 않으면 Word 기본값이 나와 자소가 깨져 보이는 환경이 있다.
FONT = '맑은 고딕'

MUTED = RGBColor(0x64, 0x74, 0x8B)
WARN = RGBColor(0xB4, 0x53, 0x09)


def _style(run, size=10.5, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # 한글은 eastasia 글꼴을 따로 지정해야 적용된다.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia',
        FONT)
    return run


def _para(doc, text, **kw):
    p = doc.add_paragraph()
    _style(p.add_run(text), **kw)
    return p


def _write_table(doc, block):
    head, rows = block.get('head') or [], block.get('rows') or []
    if not head:
        return
    table = doc.add_table(rows=1, cols=len(head))
    table.style = 'Table Grid'
    for cell, label in zip(table.rows[0].cells, head):
        cell.text = ''
        _style(cell.paragraphs[0].add_run(str(label)), size=9, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = ''
            _style(cell.paragraphs[0].add_run(str(value)), size=9)


def _write_list(doc, block):
    for item in block.get('items') or []:
        p = doc.add_paragraph(style='List Bullet')
        _style(p.add_run(item.get('title') or ''), size=10.5)
        if item.get('tag'):
            _style(p.add_run(f"  [{item['tag']}]"), size=9, color=MUTED)
        detail = item.get('detail')
        if detail:
            # 근거 줄이 제목과 같은 무게로 보이면 목록이 안 읽힌다.
            for line in str(detail).split('\n'):
                d = doc.add_paragraph()
                d.paragraph_format.left_indent = Pt(24)
                _style(d.add_run(line), size=9, color=MUTED)


def write_document(stream, plan, view):
    """조립된 기획서를 Word 파일로 쓴다. stream 은 열린 바이너리 스트림."""
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style(title.add_run(f'{plan.year}년 디지털 트윈 전략 기획서'),
           size=20, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if view.get('status') == 'confirmed' and view.get('confirmedAt'):
        _style(sub.add_run(f"확정 {view['confirmedAt'][:10]}"),
               size=9, color=MUTED)
    else:
        # ⚠️ 초안임을 문서에 박는다. 초안이 승인본처럼 돌아다니는 것을 막는
        #    유일한 방법은 문서 자체가 그렇게 말하는 것이다.
        _style(sub.add_run('초안 — 확정 전입니다. 내용은 진단·이슈가 바뀌면 '
                           '따라 바뀝니다.'), size=9, color=WARN)

    for section in view.get('sections') or []:
        if not section.get('included'):
            continue
        doc.add_page_break()
        _para(doc, section['title'], size=15, bold=True)

        if section.get('empty'):
            # 빈 장을 빼지 않는다. 장이 없으면 읽는 사람은 그 장이 없다는 것도
            # 모른 채 "검토했겠거니" 한다.
            _style(_para(doc, '').add_run(
                '아직 비어 있습니다. ' + (section.get('hint') or '')),
                size=10, color=WARN)
            continue

        for block in section['blocks']:
            kind = block.get('type')
            if kind == 'text':
                _para(doc, block.get('text') or '')
            elif kind == 'note':
                _para(doc, block.get('text') or '', size=9, color=MUTED)
            elif kind == 'table':
                _write_table(doc, block)
                doc.add_paragraph()
            elif kind == 'list':
                _write_list(doc, block)

    doc.save(stream)
    return stream
